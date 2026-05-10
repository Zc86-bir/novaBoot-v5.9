from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(31, 78, 121)
TEXT = RGBColor(51, 51, 51)
MUTED = RGBColor(102, 102, 102)
LIGHT = "D9E2F3"


def set_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=10.5, bold=False, color=TEXT, font="Microsoft YaHei"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def make_paragraph(cell_or_doc, text="", size=10.5, bold=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2, line=1.15):
    p = cell_or_doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        style_run(r, size=size, bold=bold, color=color)
    return p


def add_labeled_line(cell_or_doc, label, value, size=10.2, value_color=TEXT):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r1 = p.add_run(label)
    style_run(r1, size=size, bold=True, color=ACCENT)
    r2 = p.add_run(value)
    style_run(r2, size=size, color=value_color)
    return p


def add_section_title(cell_or_doc, title):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(title)
    style_run(r, size=11.5, bold=True, color=ACCENT)
    return p


def add_bullet(cell_or_doc, text, size=10.1):
    p = cell_or_doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.32)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.12
    r1 = p.add_run("• ")
    style_run(r1, size=size, bold=True, color=ACCENT)
    r2 = p.add_run(text)
    style_run(r2, size=size, color=TEXT)
    return p


doc = Document()
section = doc.sections[0]
set_page(section)
section.start_type = WD_SECTION.CONTINUOUS

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Cm(6.1)
table.columns[1].width = Cm(12.8)

left = table.cell(0, 0)
right = table.cell(0, 1)

for cell in (left, right):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_padding(cell, top=110, bottom=110, left=140, right=140)
    set_cell_border(cell, color="D9E2F3", size="4")

set_cell_shading(left, LIGHT)

name = make_paragraph(left, "邓杰元", size=18, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line=1)
subtitle = make_paragraph(left, "应届生 / AI与软件方向", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, line=1)

add_section_title(left, "基本信息")
add_labeled_line(left, "电话：", "13795983705")
add_labeled_line(left, "邮箱：", "3314124537@qq.com")
add_labeled_line(left, "意向城市：", "成都")
add_labeled_line(left, "籍贯：", "四川渠县")
add_labeled_line(left, "政治面貌：", "共青团员")
add_labeled_line(left, "英语：", "CET-6 558分")

add_section_title(left, "技能概览")
add_bullet(left, "编程语言：Java、Python、C、Verilog HDL")
add_bullet(left, "算法方向：目标检测、目标跟踪、多模态视觉理解")
add_bullet(left, "工程能力：需求分析、方案设计、模块联调、功能测试")
add_bullet(left, "其他工具：Matlab")

add_section_title(left, "荣誉奖励")
add_bullet(left, "校“三好学生”（2021-2022学年）")
add_bullet(left, "四川省大学生工程训练综合能力竞赛团体三等奖")
add_bullet(left, "泰迪杯数据分析大赛团体三等奖")
add_bullet(left, "创研杯全国大学生英语翻译竞赛全国二等奖")
add_bullet(left, "全国高校创新英语挑战活动综合能力赛全国二等奖")

add_section_title(left, "个人优势")
add_bullet(left, "学习适应快，能较快进入新技术和新业务场景")
add_bullet(left, "沟通协作顺畅，具备志愿服务和团队配合经历")
add_bullet(left, "面对问题能够持续跟进并主动寻找解决方案")

header = right.add_paragraph()
header.paragraph_format.space_after = Pt(6)
header.paragraph_format.line_spacing = 1
r = header.add_run("求职简历")
style_run(r, size=18, bold=True, color=ACCENT)

summary = make_paragraph(
    right,
    "具备电子信息与软件结合背景，关注计算机视觉与工程落地。拥有毕业设计、项目实习与竞赛经历，能够完成需求分析、方案设计、开发联调与测试验证，适合应届生 Java / AI 应用 / 软件开发相关岗位。",
    size=10.5,
    color=TEXT,
    space_after=6,
    line=1.2,
)

add_section_title(right, "教育背景")
edu1 = right.add_paragraph()
edu1.paragraph_format.space_after = Pt(1)
edu1.paragraph_format.line_spacing = 1.1
r = edu1.add_run("西华大学")
style_run(r, size=11.2, bold=True, color=TEXT)
r = edu1.add_run("  本科")
style_run(r, size=10.2, color=MUTED)

edu2 = right.add_paragraph()
edu2.paragraph_format.space_after = Pt(1)
edu2.paragraph_format.line_spacing = 1.1
r = edu2.add_run("相关方向：")
style_run(r, size=10.2, bold=True, color=ACCENT)
r = edu2.add_run("电子信息 / 通信类课程背景（依据原简历课程信息整理）")
style_run(r, size=10.2, color=TEXT)

course_text = "核心课程：C语言、Python、电路分析基础、模拟电子技术、数字电子技术、信号与系统、数字信号处理、通信原理、数据与计算机通信、无线通信原理与移动网络、光纤通信、Java、Matlab、Verilog HDL。"
make_paragraph(right, course_text, size=9.7, color=MUTED, space_after=5, line=1.15)

add_section_title(right, "项目经历")
proj_h = right.add_paragraph()
proj_h.paragraph_format.space_after = Pt(2)
proj_h.paragraph_format.line_spacing = 1.1
r = proj_h.add_run("基于AI视频图像目标检测")
style_run(r, size=11.1, bold=True, color=TEXT)
r = proj_h.add_run("  |  毕业设计  |  2024.11 - 2025.05")
style_run(r, size=9.8, color=MUTED)
add_bullet(right, "围绕跨模态智能监控场景，构建“目标检测 + 多目标跟踪 + 行为语义理解”的完整分析链路。", size=9.9)
add_bullet(right, "基于 YOLOv11 引入 CGA 机制进行模型改进，提升复杂场景下异常事件识别能力。", size=9.9)
add_bullet(right, "集成 ByteTrack 多目标跟踪算法与 Qwen2.5-VL 多模态模型，完成从检测到语义理解的系统联动。", size=9.9)
add_bullet(right, "负责模型改进与训练、跟踪模块设计实现，以及系统功能与性能测试。", size=9.9)

add_section_title(right, "实习经历")
job_h = right.add_paragraph()
job_h.paragraph_format.space_after = Pt(2)
job_h.paragraph_format.line_spacing = 1.1
r = job_h.add_run("粤嵌科技")
style_run(r, size=11.1, bold=True, color=TEXT)
r = job_h.add_run("  |  实习生  |  2024.09 - 2025.12（原简历时间）")
style_run(r, size=9.8, color=MUTED)
add_bullet(right, "参与指纹考察机设计项目实习，跟进项目需求分析并协助制定实现方案。", size=9.9)
add_bullet(right, "参与软硬件模块开发与联调，配合定位问题并优化设备性能。", size=9.9)
add_bullet(right, "在项目协作中积累了从需求理解到测试验证的完整工程实践经验。", size=9.9)

add_section_title(right, "语言与实践")
add_bullet(right, "曾担任中泰教育交流会及成都大运会语言志愿者，具备较好的口语沟通与英文资料阅读能力。", size=9.9)
add_bullet(right, "Java 基础较好，具备面向对象编程意识和常用设计模式理解。", size=9.9)

out_path = r"D:\NovaBoot-v5.9\邓杰元-简历优化版.docx"
doc.save(out_path)
print(out_path)
